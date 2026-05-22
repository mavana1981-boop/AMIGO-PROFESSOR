from flask import Blueprint, render_template, redirect, url_for, flash, request, send_file, jsonify
from flask_login import login_required, current_user
from app import db
from models.models import RegistroAvaliacao, Aluno, Turma, FraseRAv
from datetime import datetime
import io

rav_bp = Blueprint("rav", __name__, url_prefix="/rav")

TEXTO_F = """Orientações: Professor(a), ao elaborar o Registro de Avaliação do 2º Ciclo (RAv) é importante considerar a descrição do processo de aprendizagem do estudante, conforme as DIRETRIZES DE AVALIAÇÃO (2014, p. 49) que assim dispõe: "é preciso que o mesmo contenha elementos da avaliação diagnóstica observados pelo docente e/ou pelo Conselho de Classe: as aprendizagens evidenciadas e as dificuldades percebidas devem ser descritas na primeira parte do documento. Em seguida, deve-se apresentar as estratégias utilizadas ou as intervenções conduzidas para sanar tais dificuldades, bem como os resultados das intervenções e outras orientações que se fizerem necessárias para que o registro cumpra a sua função formativa". Além de considerar o Currículo em Movimento do Distrito Federal Ensino Fundamental - Anos Iniciais/Anos Finais da SEEDF, a avaliação para as aprendizagens e o Projeto Político Pedagógico da unidade escolar. No caso dos estudantes atendidos pelo Programa SuperAção, além das orientações já mencionadas, deve-se adicionar informações com ênfase nas aprendizagens alcançadas pelo estudante, em conformidade com o Caderno do Programa SuperAção - Atendimento aos Estudantes em Situação de Incompatibilidade Idade/Ano do Ensino Fundamental e com a Organização Curricular do Programa SuperAção. Recomenda-se não transcrever os Conteúdos e Ações Didático-Pedagógicas, pois esses são voltados para a turma. OBSERVAÇÕES GERAIS: O RAv – Formulário 1: Descrição do Processo de Aprendizagem do Estudante é o documento oficial da Secretaria de Estado de Educação do Distrito Federal, o qual: a) deve ser apresentado às unidades orgânicas da gestão pedagógica central e intermediária desta Secretaria, quando solicitado; b) Constitui documento de escrituração escolar que também compõe o dossiê do estudante, onde deverá anexar informações de estratégias pedagógicas do formulário de adequação curricular para estudantes com deficiência e TEA, cujo original deve acompanhá-lo em caso de transferência; c) Deve ser compartilhado com as famílias e/ou os responsáveis legais e com o próprio estudante, ao final de cada bimestre; d) Constitui fonte informativa para o trabalho pedagógico com o estudante; e) Deve ser preenchido sem emendas ou rasuras; f) O Campo "Resultado Final" deve ser preenchido apenas ao final do 4º Bimestre, marcando: f.1.) Cursando, para todos os estudantes beneficiados com a "Adequação Curricular na Temporalidade"; f.2.) Progressão Continuada, para estudantes promovidos do 1º ano para o 2º ano do 1º Bloco; estudantes promovidos do 2º ano para o 3º ano do 1º Bloco; estudantes promovidos do 4º ano para o 5º ano do 2º bloco, que não excederam aos 25% de faltas permitidas, nos termos do Regimento Escolar da Rede Pública de Ensino do Distrito Federal; e, no caso dos estudantes atendidos pelo Programa SuperAção promovidos do Grupo 2 (4º ano) para o 5º ano; f.3.) Aprovado, para os estudantes do 3º ano do 1º Bloco e para os estudantes do 5º ano do 2º Bloco que obtiveram desempenho escolar exitoso e não excederam aos 25% de faltas permitidas, nos termos do Regimento Escolar da Rede Pública de Ensino do Distrito Federal; e, no caso dos estudantes atendidos pelo Programa SuperAção, para os estudantes do Grupo 1 (3º ano) e do Grupo 3 (5º ano) que avançarem 1 ano de escolaridade. f.4.) Reprovado, para aqueles estudantes do 3º ano do 1º Bloco, do 5º ano do 2º Bloco e os atendidos no Programa SuperAção Grupo 1 (3º ano) e Grupo 3 (5º ano) que não obtiveram desempenho escolar exitoso, se for o caso, bem como para aqueles estudantes do 2º Ciclo que excederam aos 25% de faltas permitidas, nos termos do Regimento Escolar. Destaca-se que os estudantes atendidos no Programa SuperAção permanecerão matriculados nos grupos/anos de origem; f.5.) Abandono nos termos do Regimento Escolar da Rede Pública de Ensino do Distrito Federal; f.6.) Avanço das Aprendizagens - Correção de Fluxo em 2 anos, para os estudantes atendidos pelo Programa Superação e que consolidaram os objetivos de aprendizagem correspondente aos 2 anos escolares. g) O RAv – Formulário 1 deve ser assinado pelos(as) Professores(as), Coordenador(a) Pedagógico(a) e pai/mãe ou responsável legal do estudante; h) Nas turmas atendidas por mais de um(a) professor(a), como por exemplo na Educação em Tempo Integral, a elaboração deverá ser coletiva e todos e assinarão um único relatório; i) No caso dos estudantes atendidos na Rede Integradora do Plano Piloto, os relatórios emitidos pela Escola Parque deverão ser anexados ao RAv - Formulário 1: Descrição do Processo de Aprendizagem do Estudante, ao final de cada bimestre."""

TEXTO_G = 'O presente formulário é composto por sete itens (de "A" a "G") os quais, em nenhuma hipótese poderão ser excluídos, considerando o caráter institucional do documento, haja vista que sua modificação retira a fé pública nele depositada.\nBom trabalho!\nDiretoria do Ensino Fundamental/UNIGEEB/SUBEB'


def get_alunos():
    turmas = Turma.query.filter_by(professor_id=current_user.id).all()
    alunos = []
    for t in turmas:
        alunos.extend(sorted(t.alunos, key=lambda a: a.nome))
    return sorted(alunos, key=lambda a: a.nome)


@rav_bp.route("/")
@login_required
def index():
    turmas   = Turma.query.filter_by(professor_id=current_user.id).all()
    turma_id = request.args.get("turma_id", type=int)
    ravs     = RegistroAvaliacao.query.filter_by(professor_id=current_user.id)\
                                      .order_by(RegistroAvaliacao.criado_em.desc()).all()
    if turma_id:
        ravs = [r for r in ravs if r.aluno.turma_id == turma_id]
    return render_template("rav/index.html", ravs=ravs, turmas=turmas, turma_id=turma_id)


@rav_bp.route("/novo", methods=["GET", "POST"])
@login_required
def novo():
    alunos   = get_alunos()
    aluno_id = request.args.get("aluno_id", type=int)
    frases_custom = FraseRAv.query.filter_by(professor_id=current_user.id).all()
    if request.method == "POST":
        rav = RegistroAvaliacao(
            aluno_id=request.form.get("aluno_id", type=int),
            professor_id=current_user.id,
            bimestre=request.form.get("bimestre", type=int),
            ano_letivo=request.form.get("ano_letivo", 2026, type=int),
            total_dias=request.form.get("total_dias", 48, type=int),
            total_faltas=request.form.get("total_faltas", 0, type=int),
            data_inicio_bim=request.form.get("data_inicio_bim", "12 de fevereiro"),
            data_fim_bim=request.form.get("data_fim_bim", "29 de maio de 2026"),
            comportamento=request.form.get("comportamento", ""),
            linguagem_leitura=request.form.get("linguagem_leitura", ""),
            linguagem_escrita=request.form.get("linguagem_escrita", ""),
            gramatica=request.form.get("gramatica", ""),
            matematica=request.form.get("matematica", ""),
            ciencias=request.form.get("ciencias", ""),
            historia=request.form.get("historia", ""),
            geografia=request.form.get("geografia", ""),
            artes=request.form.get("artes", ""),
            educacao_fisica=request.form.get("educacao_fisica", ""),
            sintese=request.form.get("sintese", ""),
            perspectiva=request.form.get("perspectiva", ""),
            resultado_final=request.form.get("resultado_final", ""),
        )
        db.session.add(rav)
        db.session.commit()
        flash("RAv salvo!", "success")
        return redirect(url_for("rav.index"))
    return render_template("rav/form.html", alunos=alunos, aluno_id=aluno_id, rav=None, frases_custom=frases_custom)


@rav_bp.route("/editar/<int:id>", methods=["GET", "POST"])
@login_required
def editar(id):
    rav    = RegistroAvaliacao.query.filter_by(id=id, professor_id=current_user.id).first_or_404()
    alunos = get_alunos()
    frases_custom = FraseRAv.query.filter_by(professor_id=current_user.id).all()
    if request.method == "POST":
        rav.bimestre=request.form.get("bimestre", type=int)
        rav.ano_letivo=request.form.get("ano_letivo", 2026, type=int)
        rav.total_dias=request.form.get("total_dias", 48, type=int)
        rav.total_faltas=request.form.get("total_faltas", 0, type=int)
        rav.data_inicio_bim=request.form.get("data_inicio_bim", "")
        rav.data_fim_bim=request.form.get("data_fim_bim", "")
        rav.comportamento=request.form.get("comportamento", "")
        rav.linguagem_leitura=request.form.get("linguagem_leitura", "")
        rav.linguagem_escrita=request.form.get("linguagem_escrita", "")
        rav.gramatica=request.form.get("gramatica", "")
        rav.matematica=request.form.get("matematica", "")
        rav.ciencias=request.form.get("ciencias", "")
        rav.historia=request.form.get("historia", "")
        rav.geografia=request.form.get("geografia", "")
        rav.artes=request.form.get("artes", "")
        rav.educacao_fisica=request.form.get("educacao_fisica", "")
        rav.sintese=request.form.get("sintese", "")
        rav.perspectiva=request.form.get("perspectiva", "")
        rav.resultado_final=request.form.get("resultado_final", "")
        db.session.commit()
        flash("RAv atualizado!", "success")
        return redirect(url_for("rav.index"))
    return render_template("rav/form.html", alunos=alunos, aluno_id=rav.aluno_id, rav=rav, frases_custom=frases_custom)


@rav_bp.route("/excluir/<int:id>", methods=["POST"])
@login_required
def excluir(id):
    rav = RegistroAvaliacao.query.filter_by(id=id, professor_id=current_user.id).first_or_404()
    db.session.delete(rav)
    db.session.commit()
    flash("RAv excluído.", "info")
    return redirect(url_for("rav.index"))


# ── Frases customizadas ───────────────────────────────────────────────────────

@rav_bp.route("/frases/adicionar", methods=["POST"])
@login_required
def adicionar_frase():
    area  = request.form.get("area", "").strip()
    texto = request.form.get("texto", "").strip()
    if area and texto:
        db.session.add(FraseRAv(professor_id=current_user.id, area=area, texto=texto))
        db.session.commit()
        flash("Frase salva na sua biblioteca!", "success")
    return redirect(request.referrer or url_for("rav.index"))


@rav_bp.route("/frases/excluir/<int:id>", methods=["POST"])
@login_required
def excluir_frase(id):
    f = FraseRAv.query.filter_by(id=id, professor_id=current_user.id).first_or_404()
    db.session.delete(f)
    db.session.commit()
    return jsonify({"ok": True})


# ── Geração do .docx ─────────────────────────────────────────────────────────

@rav_bp.route("/gerar/<int:id>")
@login_required
def gerar(id):
    rav       = RegistroAvaliacao.query.filter_by(id=id, professor_id=current_user.id).first_or_404()
    try:
        buffer = gerar_docx(rav, rav.aluno, rav.aluno.turma, rav.professor)
        nome = f"RAv_{rav.aluno.nome.replace(' ','_')}_{rav.bimestre}Bim_{rav.ano_letivo}.docx"
        return send_file(buffer, as_attachment=True, download_name=nome,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        flash(f"Erro ao gerar documento: {e}", "danger")
        return redirect(url_for("rav.index"))


def gerar_docx(rav, aluno, turma, professor):
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bim     = rav.bimestre or 1
    bim_lbl = f"{bim}º Bimestre"
    bloco1  = bim in [1, 2]

    def xou(c): return "x" if c else " "
    def_nao=xou(not aluno.apresenta_deficiencia); def_sim=xou(aluno.apresenta_deficiencia)
    adeq_nao=xou(not aluno.adequacao_curricular); adeq_sim=xou(aluno.adequacao_curricular)
    temp_nao=xou(not aluno.indicado_temporalidade); temp_sim=xou(aluno.indicado_temporalidade)
    sala_nao=xou(not aluno.sala_recursos); sala_sim=xou(aluno.sala_recursos)
    sup_nao=xou(not aluno.programa_superacao); sup_sim=xou(aluno.programa_superacao)
    atend=aluno.tipo_atendimento or ""
    cl_x=xou(atend=="classe_comum"); sa_x=xou(atend=="superacao"); sr_x=xou(atend=="superacao_reduzida")
    org=aluno.org_curricular_superacao or "nao"
    org_n=xou(org=="nao"); org_s=xou(org=="sim"); org_p=xou(org=="parcialmente")
    res=rav.resultado_final or ""
    r_cur=xou(res=="cursando"); r_pro=xou(res=="progressao"); r_ava=xou(res=="avanco")
    r_apr=xou(res=="aprovado"); r_rep=xou(res=="reprovado"); r_aba=xou(res=="abandono")
    b1x=xou(bloco1); b2x=xou(not bloco1)

    regional  = (professor.regional or "").upper()
    escola    = professor.escola or ""
    ano_turma = turma.serie or f"{turma.ano}º Ano"
    prof_nome = professor.nome.upper()

    partes = [
        rav.comportamento or "", rav.linguagem_leitura or "", rav.linguagem_escrita or "",
        rav.gramatica or "", rav.matematica or "", rav.ciencias or "",
        rav.historia or "", rav.geografia or "", rav.artes or "",
        rav.educacao_fisica or "", rav.sintese or "", rav.perspectiva or "",
    ]
    texto_b = " ".join(p.strip() for p in partes if p.strip())

    # ── Documento ────────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Emu(1637190)   # 2.86 cm
    sec.bottom_margin = Emu(660400)    # 1.16 cm
    sec.left_margin   = Emu(714375)    # 1.25 cm
    sec.right_margin  = Emu(142875)    # 0.25 cm

    def run(p, text, size=10, bold=False, italic=False):
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold  = bold
        r.italic = italic
        return r

    def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=10, bold=False, before=0, after=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        if text:
            run(p, text, size, bold)
        return p

    def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar = OxmlElement('w:tcMar')
        for s, v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
            n = OxmlElement(f'w:{s}')
            n.set(qn('w:w'), str(v)); n.set(qn('w:type'), 'dxa')
            mar.append(n)
        tcPr.append(mar)

    def shade_cell(cell, fill="D9D9D9"):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    def cell_para(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run(p, text, size, bold, italic)
        return p

    def set_row_height(row, cm_val):
        tr = row._tr; trPr = tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(int(Cm(cm_val).emu / 914.4)))
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    run(p, "REGISTRO DE AVALIAÇÃO - RAv", 12.5, True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    run(p2, "Formulário 1: Descrição do Processo de Aprendizagem do Estudante", 10, False)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(4)
    run(p3, "Ensino Fundamental (Anos Iniciais)", 10, False)

    # Largura útil em EMU: 21cm - 1.25cm - 0.25cm = 19.5cm
    W_TOTAL = Emu(Cm(19.5))
    W_LETRA_A = Emu(Cm(1.73))   # coluna da letra no bloco A (245110 EMU original)
    W_LETRA_B = Emu(Cm(1.27))   # coluna da letra nos blocos seguintes
    W_CONTEUDO_A = W_TOTAL - W_LETRA_A
    W_CONTEUDO_B = W_TOTAL - W_LETRA_B

    # ── Tabela Bloco A ─────────────────────────────────────────────────────────
    tA = doc.add_table(rows=0, cols=2)
    tA.style = "Table Grid"

    linhas_a = [
        ("A", f"Ano Letivo: {rav.ano_letivo}", False),
        ("",  f"Coordenação Regional de Ensino: {regional}", False),
        ("",  f"Unidade Escolar: {escola}", False),
        ("",  f"Bloco: ({b1x}) 1º Bloco  ({b2x}) 2º Bloco", False),
        ("",  f"Ano: {ano_turma}     Turma: {turma.nome}     Turno: ( )Matutino  (x)Vespertino  ( )Integral", False),
        ("",  f"Professor (a) Regente da turma: {prof_nome}", False),
        ("",  f"Professor(a): ", False),
        ("",  f"Professor(a): ", False),
        ("",  f"Professor(a): ", False),
        ("",  f"Estudante: {aluno.nome.upper()}", False),
        ("",  f"Apresenta Deficiência ou TEA? ({def_nao}) não  ({def_sim}) sim", False),
        ("",  f"Houve adequação curricular? ({adeq_nao}) não  ({adeq_sim}) sim", False),
        ("",  f"Estudante indicado para temporalidade? ({temp_nao}) não  ({temp_sim}) sim", False),
        ("",  f"Está sendo atendido em Sala de Recursos? ({sala_nao}) não  ({sala_sim}) sim", False),
        ("",  f'Estudante do Programa SuperAção "setado" no Sistema de Gestão i-Educar? ({sup_nao}) não  ({sup_sim}) sim\nAtendimento:\n({cl_x}) Classe Comum com atendimento personalizado\n({sa_x}) Turma SuperAção\n({sr_x}) Turma SuperAção Reduzida', False),
        ("",  f"Foi aplicada a Organização Curricular específica do Programa Superação?\n({org_n}) não  ({org_s}) sim  ({org_p}) parcialmente", False),
        ("",  f"{bim_lbl}     Total de dias letivos: {rav.total_dias}     Total de Faltas: {rav.total_faltas}", True),
    ]

    for letra, texto, bold in linhas_a:
        row = tA.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = W_LETRA_A; c1.width = W_CONTEUDO_A
        set_cell_margins(c0, 40, 40, 60, 60)
        set_cell_margins(c1, 40, 40, 80, 80)
        c0._tc.get_or_add_tcPr()
        if letra == "A":
            shade_cell(c0)
            cell_para(c0, "A", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
        else:
            c0.text = ""

        # Handle multiline
        c1.text = ""
        linhas = texto.split("\n")
        for i, linha in enumerate(linhas):
            if i == 0:
                p = c1.paragraphs[0]
            else:
                p = c1.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run(p, linha, 10, bold)

    # Bloco B — texto narrativo
    row_b = tA.add_row()
    c0b, c1b = row_b.cells[0], row_b.cells[1]
    c0b.width = W_LETRA_A; c1b.width = W_CONTEUDO_A
    set_cell_margins(c0b, 40, 40, 60, 60)
    set_cell_margins(c1b, 60, 60, 80, 80)
    shade_cell(c0b)
    cell_para(c0b, "B", 10, True, WD_ALIGN_PARAGRAPH.CENTER)

    c1b.text = ""
    pb = c1b.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after  = Pt(0)
    # Line spacing 1.0
    pPr = pb._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)
    run(pb, texto_b, 10)

    # ── Tabela C D E ──────────────────────────────────────────────────────────
    tCDE = doc.add_table(rows=0, cols=2)
    tCDE.style = "Table Grid"

    # Linha vazia de continuação (overflow do B)
    row_cont = tCDE.add_row()
    c0c, c1c = row_cont.cells[0], row_cont.cells[1]
    c0c.width = W_LETRA_B; c1c.width = W_CONTEUDO_B
    set_cell_margins(c0c); set_cell_margins(c1c, 60, 60, 80, 80)
    c0c.text = ""; c1c.text = ""

    # C — Local/Data
    row_c = tCDE.add_row()
    c0, c1 = row_c.cells[0], row_c.cells[1]
    c0.width = W_LETRA_B; c1.width = W_CONTEUDO_B
    set_cell_margins(c0); set_cell_margins(c1)
    shade_cell(c0)
    cell_para(c0, "C", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(c1, f"Local/Data: Brasília/DF, {rav.data_fim_bim or '29 de maio de 2026'}.", 10, True)

    # D — Assinaturas (3 linhas, 2 colunas)
    assinaturas = [
        (f"Assinatura/Matrícula da Professora Regente", "Assinatura/Matrícula do(a) Professor(a)"),
        ("Assinatura/Matrícula do(a) Professor(a)", "Assinatura/Matrícula do(a) Coordenador(a) Pedagógico"),
        ("Assinatura/Matrícula do(a) Professor(a)", "Assinatura do(a) Pai/Mãe ou Responsável Legal"),
    ]
    for i, (esq, dir_) in enumerate(assinaturas):
        row_d = tCDE.add_row()
        set_row_height(row_d, 1.5)
        c_letra = row_d.cells[0]
        # Mescla coluna D verticalmente usando spans não disponível facilmente — usa letra só na 1ª
        c_letra.width = W_LETRA_B
        set_cell_margins(c_letra)
        if i == 0:
            shade_cell(c_letra)
            cell_para(c_letra, "D", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
        else:
            c_letra.text = ""

        # Divide coluna de conteúdo em 2 sub-colunas via tabela interna
        c_cont = row_d.cells[1]
        c_cont.width = W_CONTEUDO_B
        set_cell_margins(c_cont, 0, 0, 0, 0)
        c_cont.text = ""
        # Tabela interna 1x2
        inner = c_cont.add_table(rows=1, cols=2)
        inner.style = "Table Grid"
        w_half = W_CONTEUDO_B // 2
        ic0, ic1 = inner.rows[0].cells[0], inner.rows[0].cells[1]
        ic0.width = w_half; ic1.width = w_half
        set_cell_margins(ic0, 200, 80, 80, 80); set_cell_margins(ic1, 200, 80, 80, 80)
        cell_para(ic0, esq, 9)
        cell_para(ic1, dir_, 9)

    # E — Resultado Final
    row_e = tCDE.add_row()
    c0e, c1e = row_e.cells[0], row_e.cells[1]
    c0e.width = W_LETRA_B; c1e.width = W_CONTEUDO_B
    set_cell_margins(c0e); set_cell_margins(c1e, 80, 80, 80, 80)
    shade_cell(c0e)
    cell_para(c0e, "E", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    c1e.text = ""
    pe1 = c1e.paragraphs[0]
    pe1.paragraph_format.space_before = Pt(0); pe1.paragraph_format.space_after = Pt(2)
    run(pe1, "Resultado Final (Preencher somente ao final do 4º bimestre)", 10, True)
    for linha in [
        f"({r_cur}) Cursando  ({r_pro}) Progressão Continuada  ({r_ava}) Avanço das Aprendizagens - Correção de Fluxo",
        f"({r_apr}) Aprovado  ({r_rep}) Reprovado  ({r_aba}) Abandono",
    ]:
        pe = c1e.add_paragraph()
        pe.paragraph_format.space_before = Pt(0); pe.paragraph_format.space_after = Pt(0)
        run(pe, linha, 10)

    # ── Tabela F ───────────────────────────────────────────────────────────────
    tF = doc.add_table(rows=1, cols=2)
    tF.style = "Table Grid"
    c0f, c1f = tF.rows[0].cells[0], tF.rows[0].cells[1]
    c0f.width = W_LETRA_B; c1f.width = W_CONTEUDO_B
    set_cell_margins(c0f); set_cell_margins(c1f, 60, 60, 80, 80)
    shade_cell(c0f)
    cell_para(c0f, "F", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    c1f.text = ""
    pf = c1f.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.paragraph_format.space_before = Pt(0); pf.paragraph_format.space_after = Pt(0)
    run(pf, TEXTO_F, 9)

    # ── Tabela G ───────────────────────────────────────────────────────────────
    tG = doc.add_table(rows=1, cols=2)
    tG.style = "Table Grid"
    c0g, c1g = tG.rows[0].cells[0], tG.rows[0].cells[1]
    c0g.width = W_LETRA_B; c1g.width = W_CONTEUDO_B
    set_cell_margins(c0g); set_cell_margins(c1g, 60, 60, 80, 80)
    shade_cell(c0g)
    cell_para(c0g, "G", 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    c1g.text = ""
    pg = c1g.paragraphs[0]
    pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pg.paragraph_format.space_before = Pt(0); pg.paragraph_format.space_after = Pt(0)
    run(pg, TEXTO_G, 9)

    # Rodapé versão
    pv = doc.add_paragraph()
    pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pv.paragraph_format.space_before = Pt(2)
    run(pv, f"RAv_versão {rav.ano_letivo}", 8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
